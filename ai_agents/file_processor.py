import os
from typing import Optional, BinaryIO
from pathlib import Path
import PyPDF2
import docx

class FileProcessor:
    """
    Utility for processing various document formats and extracting text.
    Supports: PDF, DOCX, TXT
    """
    
    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.doc', '.txt', '.md'}
    
    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str:
        """Extract text from PDF file"""
        text = []
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text.append(page.extract_text())
                
            return "\n\n".join(text)
        
        except Exception as e:
            raise ValueError(f"Failed to extract text from PDF: {e}")
    
    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(file_path)
            
            text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text.append(para.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text.append(cell.text)
            
            return "\n\n".join(text)
        
        except Exception as e:
            raise ValueError(f"Failed to extract text from DOCX: {e}")
    
    @staticmethod
    def extract_text_from_txt(file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    return file.read()
            except Exception as e:
                raise ValueError(f"Failed to read text file: {e}")
        except Exception as e:
            raise ValueError(f"Failed to extract text from TXT: {e}")
    
    @classmethod
    def extract_text(cls, file_path: str) -> str:
        """
        Extract text from any supported file type.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Extracted text content
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        extension = path.suffix.lower()
        
        if extension not in cls.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file type: {extension}. "
                f"Supported types: {', '.join(cls.SUPPORTED_EXTENSIONS)}"
            )
        
        if extension == '.pdf':
            return cls.extract_text_from_pdf(file_path)
        elif extension in {'.docx', '.doc'}:
            return cls.extract_text_from_docx(file_path)
        elif extension in {'.txt', '.md'}:
            return cls.extract_text_from_txt(file_path)
        else:
            raise ValueError(f"Handler not implemented for {extension}")
    
    @staticmethod
    def save_uploaded_file(file: BinaryIO, filename: str, upload_dir: str = "./uploads") -> str:
        """
        Save an uploaded file to disk.
        
        Args:
            file: File object
            filename: Original filename
            upload_dir: Directory to save files
            
        Returns:
            Path to saved file
        """
        os.makedirs(upload_dir, exist_ok=True)
        
        # Sanitize filename
        safe_filename = "".join(c for c in filename if c.isalnum() or c in "._- ")
        file_path = os.path.join(upload_dir, safe_filename)
        
        with open(file_path, 'wb') as f:
            f.write(file.read())
        
        return file_path
    
    @staticmethod
    def get_file_info(file_path: str) -> dict:
        """Get information about a file"""
        path = Path(file_path)
        
        return {
            "filename": path.name,
            "extension": path.suffix,
            "size_bytes": path.stat().st_size,
            "size_mb": round(path.stat().st_size / (1024 * 1024), 2),
            "is_supported": path.suffix.lower() in FileProcessor.SUPPORTED_EXTENSIONS
        }

# Example usage
if __name__ == "__main__":
    import sys
    
    if len(sys.argv) > 1:
        file_path = sys.argv[1]
        
        try:
            print(f"Processing: {file_path}")
            info = FileProcessor.get_file_info(file_path)
            print(f"File info: {info}")
            
            if info["is_supported"]:
                text = FileProcessor.extract_text(file_path)
                print(f"\nExtracted text ({len(text)} characters):")
                print("="*80)
                print(text[:1000])  # First 1000 chars
                if len(text) > 1000:
                    print(f"\n... ({len(text) - 1000} more characters)")
            else:
                print("File type not supported")
        
        except Exception as e:
            print(f"Error: {e}")
    else:
        print("Usage: python file_processor.py <file_path>")
